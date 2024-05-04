import markdown2
import re
import sys
from docx import Document
from docx.shared import Pt

def strip_tags(text):
    """Remove HTML tags from a string."""
    return re.sub('<[^<]+?>', '', text)

def add_line(document, line):
    """Add a line to the document, handling Markdown formatting."""
    par = document.add_paragraph()
    is_bold = False
    is_subscript = False
    indent = count_initial_tabs(line)

    words = line.split()
    for i, word in enumerate(words):
        if i == 0:
            # Handle bullets and headings based on the first word
            if word == '*':
                par.style = 'List Bullet' if indent == 0 else f'List Bullet {indent + 1}'
                continue
            elif word.count('#') > 0:
                par.style = f'Heading {word.count("#")}'
                continue

        clean_word = strip_tags(word)
        if word.startswith('**'):
            clean_word = clean_word[2:]
            is_bold = not is_bold
        if word.endswith('**'):
            clean_word = clean_word[:-2]

        if word.startswith('<sub>'):
            is_subscript = True

        run = par.add_run(clean_word + ' ' if i < len(words) - 1 else clean_word)
        run.bold = is_bold
        run.font.subscript = is_subscript

        if word.endswith('</sub>'):
            is_subscript = False
        if word.endswith('**'):
            is_bold = False

def count_initial_tabs(line):
    """Count the initial tabs in a line for indentation handling."""
    return len(line) - len(line.lstrip('\t'))

def convert_markdown_to_docx(input_file, output_file):
    """Convert Markdown file to a DOCX file."""
    document = Document()

    with open(input_file, 'r', encoding='utf-8') as file:
        markdown_text = file.read()

    for sentence in markdown_text.split('\n'):
        add_line(document, sentence)

    document.save(output_file)

def main():
    """Main function to handle command line arguments."""
    if len(sys.argv) != 3:
        print("Usage: python script.py input.md output.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    convert_markdown_to_docx(input_file, output_file)

if __name__ == "__main__":
    main()
